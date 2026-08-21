from __future__ import annotations

import json
from io import BytesIO

import pytest
from docx import Document
from reportlab.pdfgen import canvas

from app.services.documents import DocumentIngestionError, parse_document


def pdf_bytes(text: str) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer)
    pdf.drawString(72, 720, text)
    pdf.save()
    return buffer.getvalue()


def docx_bytes(text: str) -> bytes:
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("Functional Specification", level=1)
    doc.add_paragraph(text)
    doc.save(buffer)
    return buffer.getvalue()


@pytest.mark.parametrize(
    ("filename", "payload", "expected_type", "needle"),
    [
        ("FSD-COL-104.txt", b"Promise-to-Pay duplicate event handling", "TXT", "duplicate event"),
        ("notes.md", b"# PTP\nIdempotency guard", "MARKDOWN", "Idempotency guard"),
        ("config.json", json.dumps({"module": "PTP", "idempotent": True}).encode(), "JSON", '"idempotent": true'),
    ],
)
def test_text_document_parsers(filename: str, payload: bytes, expected_type: str, needle: str):
    parsed = parse_document(filename, payload)
    assert parsed.document_type == expected_type
    assert needle in parsed.text
    assert len(parsed.content_hash) == 64


def test_pdf_parser_extracts_real_text():
    parsed = parse_document("FSD-COL-104.pdf", pdf_bytes("Promise-to-Pay Event Processing"), declared_content_type="application/pdf")
    assert parsed.document_type == "PDF"
    assert "Promise-to-Pay Event Processing" in parsed.text


def test_docx_parser_extracts_real_text():
    parsed = parse_document(
        "FSD-COL-104.docx",
        docx_bytes("Duplicate events require an idempotency key."),
        declared_content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )
    assert parsed.document_type == "DOCX"
    assert "Duplicate events require an idempotency key." in parsed.text


def test_rejects_unsupported_and_bad_json():
    with pytest.raises(DocumentIngestionError, match="UNSUPPORTED_FILE_TYPE"):
        parse_document("malware.exe", b"x")
    with pytest.raises(DocumentIngestionError, match="JSON_PARSE_FAILED"):
        parse_document("broken.json", b"{nope")
