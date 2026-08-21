from __future__ import annotations

import hashlib
import json
import mimetypes
import re
from dataclasses import dataclass
from io import BytesIO
from pathlib import Path
from typing import BinaryIO

from docx import Document as DocxDocument
from pypdf import PdfReader

SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md", ".json"}
MIME_BY_EXTENSION = {
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".json": "application/json",
}


class DocumentIngestionError(ValueError):
    pass


@dataclass(frozen=True)
class ParsedDocument:
    title: str
    document_type: str
    mime_type: str
    text: str
    content_hash: str
    file_size: int
    original_filename: str


def safe_filename(name: str) -> str:
    base = Path(name).name.strip()
    if not base:
        raise DocumentIngestionError("INVALID_FILENAME")
    stem = re.sub(r"[^A-Za-z0-9._-]+", "_", base)
    return stem[:220]


def document_type_from_extension(extension: str) -> str:
    return {
        ".pdf": "PDF",
        ".docx": "DOCX",
        ".txt": "TXT",
        ".md": "MARKDOWN",
        ".json": "JSON",
    }[extension]


def _decode_text(data: bytes) -> str:
    try:
        return data.decode("utf-8")
    except UnicodeDecodeError:
        try:
            return data.decode("utf-8-sig")
        except UnicodeDecodeError as exc:
            raise DocumentIngestionError("TEXT_ENCODING_UNSUPPORTED") from exc


def _parse_pdf(data: bytes) -> str:
    try:
        reader = PdfReader(BytesIO(data))
        pages = [(page.extract_text() or "").strip() for page in reader.pages]
    except Exception as exc:
        raise DocumentIngestionError("PDF_PARSE_FAILED") from exc
    text = "\n\n".join(page for page in pages if page)
    if not text.strip():
        raise DocumentIngestionError("PDF_NO_EXTRACTABLE_TEXT")
    return text


def _parse_docx(data: bytes) -> str:
    try:
        doc = DocxDocument(BytesIO(data))
        blocks = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                values = [cell.text.strip() for cell in row.cells]
                if any(values):
                    blocks.append(" | ".join(values))
    except Exception as exc:
        raise DocumentIngestionError("DOCX_PARSE_FAILED") from exc
    text = "\n".join(blocks)
    if not text.strip():
        raise DocumentIngestionError("DOCX_NO_EXTRACTABLE_TEXT")
    return text


def _parse_json(data: bytes) -> str:
    raw = _decode_text(data)
    try:
        value = json.loads(raw)
    except json.JSONDecodeError as exc:
        raise DocumentIngestionError("JSON_PARSE_FAILED") from exc
    return json.dumps(value, ensure_ascii=False, indent=2, sort_keys=True)


def parse_document(filename: str, data: bytes, *, declared_content_type: str | None = None) -> ParsedDocument:
    clean_name = safe_filename(filename)
    extension = Path(clean_name).suffix.lower()
    if extension not in SUPPORTED_EXTENSIONS:
        raise DocumentIngestionError("UNSUPPORTED_FILE_TYPE")
    if not data:
        raise DocumentIngestionError("EMPTY_FILE")

    if extension == ".pdf":
        text = _parse_pdf(data)
    elif extension == ".docx":
        text = _parse_docx(data)
    elif extension == ".json":
        text = _parse_json(data)
    else:
        text = _decode_text(data)
        if not text.strip():
            raise DocumentIngestionError("EMPTY_TEXT_DOCUMENT")

    mime_type = MIME_BY_EXTENSION[extension]
    if declared_content_type and declared_content_type not in {
        "application/octet-stream",
        mime_type,
        "text/plain" if extension == ".md" else mime_type,
    }:
        # Browser MIME reporting is inconsistent, so only reject obviously conflicting values.
        guessed, _ = mimetypes.guess_type(clean_name)
        if guessed and declared_content_type != guessed:
            raise DocumentIngestionError("MIME_TYPE_MISMATCH")

    return ParsedDocument(
        title=Path(clean_name).stem,
        document_type=document_type_from_extension(extension),
        mime_type=mime_type,
        text=text.strip(),
        content_hash=hashlib.sha256(data).hexdigest(),
        file_size=len(data),
        original_filename=clean_name,
    )


def persist_file(storage_root: Path, document_id: str, filename: str, data: bytes) -> Path:
    target_dir = storage_root / document_id
    target_dir.mkdir(parents=True, exist_ok=False)
    target = target_dir / safe_filename(filename)
    target.write_bytes(data)
    return target


def knowledge_storage_health(storage_root: Path) -> tuple[str, str | None]:
    try:
        storage_root.mkdir(parents=True, exist_ok=True)
        probe = storage_root / ".creed-write-probe"
        probe.write_text("ok", encoding="utf-8")
        probe.unlink(missing_ok=True)
        return "CONNECTED", None
    except OSError as exc:
        return "UNAVAILABLE", f"{exc.__class__.__name__}: {exc}"
