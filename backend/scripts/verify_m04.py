from __future__ import annotations

import json
import os
import subprocess
import sys
import tempfile
from io import BytesIO
from pathlib import Path

from docx import Document
from fastapi.testclient import TestClient
from reportlab.pdfgen import canvas

ROOT = Path(__file__).resolve().parents[1]
PROJECT = ROOT.parent
sys.path.insert(0, str(ROOT))


def pdf_bytes(text: str) -> bytes:
    buffer = BytesIO()
    doc = canvas.Canvas(buffer)
    doc.drawString(72, 720, text)
    doc.save()
    return buffer.getvalue()


def docx_bytes(text: str) -> bytes:
    buffer = BytesIO()
    doc = Document()
    doc.add_heading("Functional Specification", level=1)
    doc.add_paragraph(text)
    doc.save(buffer)
    return buffer.getvalue()


def run() -> int:
    with tempfile.TemporaryDirectory(prefix="creed-m04-") as temp:
        tmp = Path(temp)
        db_path = tmp / "m04.db"
        uploads = tmp / "documents"
        env = os.environ.copy()
        env["DATABASE_URL"] = f"sqlite:///{db_path}"
        env["DOCUMENT_STORAGE_PATH"] = str(uploads)
        env["PYTHONPATH"] = str(ROOT)
        subprocess.run([sys.executable, "-m", "alembic", "upgrade", "head"], cwd=ROOT, env=env, check=True, stdout=subprocess.DEVNULL)

        os.environ.update({"DATABASE_URL": env["DATABASE_URL"], "DOCUMENT_STORAGE_PATH": env["DOCUMENT_STORAGE_PATH"]})
        from app.core.config import get_settings
        from app.db.session import reset_database_caches
        get_settings.cache_clear()
        reset_database_caches()
        from app.main import app
        client = TestClient(app)

        fixtures = [
            ("FSD-COL-104.pdf", pdf_bytes("Promise-to-Pay Event Processing FSD"), "application/pdf"),
            ("TC-COL-217.docx", docx_bytes("Duplicate PTP events must not change state twice."), "application/vnd.openxmlformats-officedocument.wordprocessingml.document"),
            ("CFG-MERIDIAN-PTP.txt", b"event_mode=legacy\nidempotency_guard=false", "text/plain"),
            ("notes.md", b"# PTP Notes\nOut-of-order events require review.", "text/markdown"),
            ("CFG-NOVA-PTP.json", json.dumps({"idempotency_guard": True, "module": "PTP"}).encode(), "application/json"),
        ]
        ids: list[str] = []
        for name, body, mime in fixtures:
            response = client.post("/api/v1/documents", data={"source": "LOCAL_DEMO"}, files={"file": (name, body, mime)})
            if response.status_code != 201:
                print("FAIL upload", name, response.status_code, response.text)
                return 1
            data = response.json()
            ids.append(data["id"])
            print(f"PASS {data['document_type']:<8} {name:<26} {data['char_count']:>4} chars  {data['content_hash'][:10]}…")

        listing = client.get("/api/v1/documents").json()
        if len(listing) != 5:
            print("FAIL expected 5 documents, got", len(listing))
            return 1
        detail = client.get(f"/api/v1/documents/{ids[0]}").json()
        if "Promise-to-Pay Event Processing FSD" not in detail["extracted_text"]:
            print("FAIL PDF extracted text missing")
            return 1
        stored_files = [p for p in uploads.rglob("*") if p.is_file()]
        if len(stored_files) != 5:
            print("FAIL expected 5 stored source files, got", len(stored_files))
            return 1
        bad = client.post("/api/v1/documents", files={"file": ("malware.exe", b"x", "application/octet-stream")})
        if bad.status_code != 422:
            print("FAIL unsupported format was not rejected")
            return 1
        duplicate = client.post("/api/v1/documents", data={"source":"LOCAL_DEMO"}, files={"file": fixtures[2]})
        if duplicate.status_code != 409:
            print("FAIL duplicate was not rejected")
            return 1

        subprocess.run([sys.executable, "-m", "alembic", "downgrade", "base"], cwd=ROOT, env=env, check=True, stdout=subprocess.DEVNULL)
        print("PASS: M04 real parsing, persistence, source storage, duplicate rejection and migration verified.")
        return 0


if __name__ == "__main__":
    raise SystemExit(run())
